from docx import Document
import pandas as pd
from pathlib import Path
from utils.formatters import clean_text
import logging

logger = logging.getLogger(__name__)

def extract_from_word(docx_path: str) -> pd.DataFrame:

    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {path}")

    doc = Document(path)
    data = []
    total_tabelas = len(doc.tables)
    logger.debug(f"{total_tabelas} tabela(s) detectadas no documento Word.")

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            if i == 0 and any("nome" in cell.text.lower() for cell in row.cells):
                continue

            try:
                nome = clean_text(row.cells[0].text)
                turma = clean_text(row.cells[1].text)
                responsavel = clean_text(row.cells[2].text)
                fone = clean_text(row.cells[3].text)
                if nome:
                    data.append((nome, turma, responsavel, fone))
            except IndexError:
                logger.warning(f"Linha ignorada (faltando dados): {row.cells}")
                continue

    if not data:
        logger.warning(f"Nenhum registro válido encontrado em: {path}")
        df = pd.DataFrame(columns=["Nome", "Turma", "Responsável", "Fone"])
    else:
        df = pd.DataFrame(data, columns=["Nome", "Turma", "Responsável", "Fone"])
        logger.info(f"{len(df)} registros importados do Word com sucesso.")

    return df


def export_to_excel(df: pd.DataFrame, output_path: str = "data/alunos_importados.xlsx") -> str:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    if df.empty:
        logger.warning("Tentativa de exportar planilha vazia.")
    else:
        df.to_excel(output, index=False)
        logger.info(f"Planilha salva em: {output.resolve()}")

    return str(output.resolve())
